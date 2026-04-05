import os
import django
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'fengyun.settings')
django.setup()

from inn_app.models import SigninTemplate
from django.core.files import File

def create_template():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        
    title = doc.add_heading('{{ course_title }} 簽到名冊', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('開課時間：{{ course_date }}')
    doc.add_paragraph(' ')
    
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '序號'
    hdr[1].text = '姓名'
    hdr[2].text = '性別'
    hdr[3].text = '聯絡電話'
    hdr[4].text = '報名人數'
    hdr[5].text = '本人簽名'
    
    # Jinja Row for docxtpl
    row = table.rows[1].cells
    row[0].text = '{% tr for reg in registrations %}{{ reg.index }}'
    row[1].text = '{{ reg.name }}'
    row[2].text = '{{ reg.gender }}'
    row[3].text = '{{ reg.phone }}'
    row[4].text = '{{ reg.headcount }} 人'
    row[5].text = '{% endtr %}'
    
    for r in table.rows:
        r.height = Cm(1.5)
        
    os.makedirs('media/templates', exist_ok=True)
    temp_path = 'media/templates/客棧預設公版.docx'
    doc.save(temp_path)
    
    if not SigninTemplate.objects.filter(name="系統預設公版").exists():
        with open(temp_path, 'rb') as f:
            SigninTemplate.objects.create(
                name="系統預設公版",
                is_default=True,
                file=File(f, name='default_template.docx')
            )
            print("成功建立並上傳預設範本！")

if __name__ == '__main__':
    create_template()
